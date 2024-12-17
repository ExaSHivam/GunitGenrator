import os

from flask import Flask, render_template, request, redirect, session, jsonify
from aws_s3 import generate_gunit_data_claude, generate_gunit_data_claude_class, get_excel_from_s3, \
    filter_builders_by_lob
from git import get_sha_for_path, get_file_path, fetch_file_content
from llama_claude import generate_gunit_data_llama
from dotenv import load_dotenv
from flask import send_file, render_template, url_for
from docx import Document
from io import BytesIO

# from claude_ai import generate_gunit_data_claude, generate_gunit_data_claude_llama

app = Flask(__name__)
app.secret_key = os.urandom(24)

UPLOAD_FOLDER = './static/uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER


@app.route("/")
def home():
    return render_template('index.html')


load_dotenv()
repo_owner = os.getenv('repo_owner')
repo_name = os.getenv('repo_name')


@app.route("/get_builders", methods=['POST'])
def get_builders():
    print("entered get builder function")
    object_name = 'Builders_list.xlsx'
    builder_list = get_excel_from_s3(object_name)
    print(builder_list.to_string())
    selected_lob = request.form.get('lob')
    print("selected lob in get builders", selected_lob)
    filtered_builders = filter_builders_by_lob(builder_list, selected_lob)
    print("filtered values", filtered_builders)
    return jsonify(builder=filtered_builders)


@app.route('/download_report')
def download_report():
    temp_file_path = 'gunit_report.docx'
    return send_file(
        temp_file_path,
        as_attachment=True,
        download_name='gunit_report.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route("/generate_gunit", methods=['POST'])
def generate_gunit():
    selected_lob = request.form.get('selected_lob')
    print("lob", selected_lob)
    builder = request.form.get('builder')
    print("builder", builder)
    base_method = request.form.get('base_method')
    print("base method", base_method)
    features = request.form.getlist("features")
    print("objects", features)
    class_name = request.form.get('class_name')
    print("class name:", class_name)

    if (base_method):
        print("entered if block")
        builder_output, gunit_output = generate_gunit_data_claude(selected_lob, builder, base_method, features)
    else:
        print("entered else block")
        sha_hash = get_sha_for_path(repo_owner, repo_name)
        file_path = get_file_path(repo_owner, repo_name, sha_hash, class_name)
        print("sha_hash", sha_hash)
        file_content = fetch_file_content(repo_owner, repo_name, file_path)
        print("file content", file_content)
        builder_output, gunit_output = generate_gunit_data_claude_class(selected_lob, builder, file_content, features)
    # print("Data Builder: ", response)
    # print("Gunit: ", response1)
    # Create a Word document
    doc = Document()
    doc.add_heading('Gunit', 0)
    doc.add_heading('Builder Output:', level=1)
    doc.add_paragraph(builder_output)
    doc.add_heading('Gunit Output:', level=1)
    doc.add_paragraph(gunit_output)

    # Save the document to a temporary file
    temp_file_path = 'gunit_report.docx'
    doc.save(temp_file_path)

    # Generate a download URL for the file
    download_url = url_for('download_report')

    # Render the HTML page with a download link
    return render_template(
        'index.html',
        status="Gunit Successfully Generated",
        response=builder_output,
        response1=gunit_output,
        selected_lob=selected_lob,
        builder=builder,
        download_url=download_url
    )


if __name__ == "__main__":
    app.run(host="0.0.0.0", debug=True)
